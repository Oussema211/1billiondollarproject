import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
import queue
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sounddevice as sd
import numpy as np
import json
import time

from .pipeline import AudioPipeline, load_audio
from .realtime_pipeline import RealtimeTranscriber
from .report_generator import ReportGenerator
from .semantic_disambiguator import disambiguate_speakers
from . import config


class MedicalScribeApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Qatar Clinical Medical Scribe - Real-Time AI Documentation")
        self.geometry("1440x920")
        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        self.pipeline: AudioPipeline | None = None
        self.report_gen: ReportGenerator | None = None

        # ── Recording state ──────────────────────────────────────────────────
        self.is_recording = False
        self._audio_queue: queue.Queue = queue.Queue()
        self._recording_stream = None
        self.sample_rate = 16000
        self.live_transcriber: RealtimeTranscriber | None = None

        # ── Current results ──────────────────────────────────────────────────
        self.current_report = None
        self.current_transcript: list = []
        self.current_attending = ""
        self.current_patient = ""
        self.current_language = "en"

        self._build_sidebar()
        self._build_content()
        self._show_tab("process")

        # Load models in background so GUI opens instantly
        threading.Thread(target=self._init_backend, daemon=True).start()

    # ══════════════════════════════════════════════════════════════════════════
    # Sidebar
    # ══════════════════════════════════════════════════════════════════════════

    def _build_sidebar(self):
        sb = ctk.CTkFrame(self, width=230, corner_radius=0)
        sb.pack(side="left", fill="y", padx=0, pady=0)

        ctk.CTkLabel(
            sb, text="Clinical Scribe AI",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(pady=(20, 2))

        ctk.CTkLabel(
            sb, text="State of Qatar • قطر",
            font=ctk.CTkFont(size=12, weight="bold"), text_color="#8A1538"  # Maroon (Qatar national color)
        ).pack(pady=(0, 2))

        ctk.CTkLabel(
            sb, text="Multilingual Offline Scribe",
            font=ctk.CTkFont(size=11), text_color="gray"
        ).pack(pady=(0, 15))

        self.nav_buttons = {}
        for key, label in [
            ("process", "🩺 Consultation Note"),
            ("enroll",  "🎙 Doctor Voice ID"),
            ("reports", "📁 Patient Records"),
        ]:
            btn = ctk.CTkButton(sb, text=label, command=lambda k=key: self._show_tab(k), anchor="w")
            btn.pack(pady=6, padx=15, fill="x")
            self.nav_buttons[key] = btn

        self.status = ctk.CTkLabel(sb, text="Initializing AI models...", wraplength=190, font=ctk.CTkFont(size=11))
        self.status.pack(side="bottom", pady=20)

    # ══════════════════════════════════════════════════════════════════════════
    # Content Area
    # ══════════════════════════════════════════════════════════════════════════

    def _build_content(self):
        self.content = ctk.CTkFrame(self, corner_radius=0)
        self.content.pack(side="right", fill="both", expand=True)

        self._build_process_tab()
        self._build_enroll_tab()
        self._build_reports_tab()

    # ── Process tab ──────────────────────────────────────────────────────────

    def _build_process_tab(self):
        self.tab_process = ctk.CTkFrame(self.content)

        # ── Top Consultation Controls ─────────────────────────────────────────
        ctrl = ctk.CTkFrame(self.tab_process)
        ctrl.grid(row=0, column=0, columnspan=2, sticky="ew", padx=15, pady=(10, 5))

        ctk.CTkLabel(ctrl, text="Patient Name:").grid(row=0, column=0, padx=(10, 5), pady=10, sticky="w")
        self.entry_patient = ctk.CTkEntry(ctrl, width=200, placeholder_text="e.g. Nasser Al-Kuwari")
        self.entry_patient.grid(row=0, column=1, padx=5, pady=10, sticky="w")

        ctk.CTkLabel(ctrl, text="Doctor Name:").grid(row=0, column=2, padx=(15, 5), pady=10, sticky="w")
        self.entry_doctor = ctk.CTkEntry(ctrl, width=180, placeholder_text="e.g. Dr. Ahmed")
        self.entry_doctor.grid(row=0, column=3, padx=5, pady=10, sticky="w")

        ctk.CTkLabel(ctrl, text="Report Language:").grid(row=0, column=4, padx=(15, 5), pady=10, sticky="w")
        self.combo_language = ctk.CTkOptionMenu(
            ctrl,
            values=["Medical English (Standard / HMC)", "Arabic (العربية - التقرير الطبي)"],
            width=230,
        )
        self.combo_language.set("Medical English (Standard / HMC)")
        self.combo_language.grid(row=0, column=5, padx=5, pady=10, sticky="w")

        self.btn_record = ctk.CTkButton(
            ctrl, text="🎙 Start Consultation",
            command=self._toggle_recording, fg_color="#8A1538", hover_color="#6B102C", width=180, height=36,
            font=ctk.CTkFont(weight="bold")
        )
        self.btn_record.grid(row=0, column=6, padx=15)

        self.recording_status = ctk.CTkLabel(
            ctrl,
            text="Multilingual STT: Speaks Arabic, English, Hindi, Urdu, Tagalog, Malayalam, French → Instant Clinical Note",
            text_color="gray", font=ctk.CTkFont(size=12)
        )
        self.recording_status.grid(row=1, column=0, columnspan=7, pady=(0, 5))

        # ── Audio File Upload Bar ─────────────────────────────────────────────
        file_frame = ctk.CTkFrame(self.tab_process)
        file_frame.grid(row=1, column=0, columnspan=2, sticky="ew", padx=15, pady=5)

        ctk.CTkLabel(file_frame, text="Or load consultation recording:").pack(side="left", padx=10, pady=8)
        self.entry_audio = ctk.CTkEntry(file_frame, width=460, placeholder_text="Browse recorded audio (WAV, MP3, M4A, FLAC)...")
        self.entry_audio.pack(side="left", padx=5, pady=8)
        ctk.CTkButton(file_frame, text="Browse...", command=self._browse_audio, width=90).pack(side="left", padx=5)

        self.chk_fast_mode = ctk.CTkCheckBox(file_frame, text="⚡ Fast Mode", onvalue=1, offvalue=0)
        self.chk_fast_mode.select()
        self.chk_fast_mode.pack(side="left", padx=10)

        ctk.CTkButton(
            file_frame, text="▶ Generate Report",
            command=self._run_file_pipeline, fg_color="#27ae60", hover_color="#219653", width=140
        ).pack(side="left", padx=10)

        # ── Progress Bar ──────────────────────────────────────────────────────
        self.progress = ctk.CTkProgressBar(self.tab_process, mode="indeterminate", height=6)
        self.progress.grid(row=2, column=0, columnspan=2, sticky="ew", padx=15, pady=3)
        self.progress.stop()
        self.progress.set(0)

        # ── Main Medical Report Document View ─────────────────────────────────
        doc_header_frame = ctk.CTkFrame(self.tab_process, fg_color="transparent")
        doc_header_frame.grid(row=3, column=0, columnspan=2, sticky="ew", padx=15, pady=(5, 0))

        ctk.CTkLabel(
            doc_header_frame, text="CLINICAL CONSULTATION NOTE",
            font=ctk.CTkFont(size=15, weight="bold")
        ).pack(side="left")

        # Action Buttons
        self.btn_transcript_toggle = ctk.CTkButton(
            doc_header_frame, text="💬 View Dialogue Transcript",
            command=self._toggle_transcript_dialog, width=170, fg_color="#4b6584", hover_color="#384c63"
        )
        self.btn_transcript_toggle.pack(side="right", padx=5)

        ctk.CTkButton(
            doc_header_frame, text="📄 Export Word Document (.docx)",
            command=self._export_docx, width=220, fg_color="#2980b9", hover_color="#1f618d"
        ).pack(side="right", padx=5)

        ctk.CTkButton(
            doc_header_frame, text="📋 Copy Report",
            command=self._copy_report, width=120, fg_color="#2c3e50", hover_color="#1a252f"
        ).pack(side="right", padx=5)

        self.report_box = ctk.CTkTextbox(
            self.tab_process, font=ctk.CTkFont(family="Consolas", size=13),
            wrap="word", corner_radius=6
        )
        self.report_box.grid(row=4, column=0, columnspan=2, sticky="nsew", padx=15, pady=5)

        # Initial Welcome / Placeholder Text
        self._set_placeholder_report()

        # Mini Status Log Bar
        self.log_label = ctk.CTkLabel(self.tab_process, text="Ready for consultation.", anchor="w", text_color="gray", font=ctk.CTkFont(size=11))
        self.log_label.grid(row=5, column=0, columnspan=2, sticky="ew", padx=15, pady=2)

        # Layout Weights
        self.tab_process.grid_columnconfigure(0, weight=1)
        self.tab_process.grid_rowconfigure(4, weight=1)

    # ── Enroll tab ───────────────────────────────────────────────────────────

    def _build_enroll_tab(self):
        self.tab_enroll = ctk.CTkFrame(self.content)
        ctk.CTkLabel(
            self.tab_enroll, text="Doctor Voice Profile Enrollment",
            font=ctk.CTkFont(size=20, weight="bold")
        ).grid(row=0, column=0, columnspan=3, padx=20, pady=20, sticky="w")

        ctk.CTkLabel(
            self.tab_enroll,
            text="Enrolling your voice sample creates a local acoustic embedding so the system automatically recognizes you as the Attending Doctor across consultations.",
            wraplength=800, text_color="gray"
        ).grid(row=1, column=0, columnspan=3, padx=20, pady=(0, 20), sticky="w")

        ctk.CTkLabel(self.tab_enroll, text="Physician Name:").grid(row=2, column=0, padx=20, pady=10, sticky="w")
        self.entry_doc_name = ctk.CTkEntry(self.tab_enroll, width=320, placeholder_text="e.g. Dr. Ahmed Al-Malki")
        self.entry_doc_name.grid(row=2, column=1, padx=10, sticky="w")

        ctk.CTkButton(
            self.tab_enroll, text="Browse Audio Sample...",
            command=self._browse_enroll, width=180
        ).grid(row=2, column=2, padx=10, sticky="w")

        ctk.CTkButton(
            self.tab_enroll, text="Enroll Voice Profile",
            command=self._run_enroll, fg_color="#27ae60", hover_color="#219653", width=200, height=36
        ).grid(row=3, column=1, pady=20, sticky="w")

        self.enroll_log = ctk.CTkTextbox(self.tab_enroll, height=350)
        self.enroll_log.grid(row=4, column=0, columnspan=3, sticky="nsew", padx=20, pady=10)
        self.tab_enroll.grid_columnconfigure(1, weight=1)
        self.tab_enroll.grid_rowconfigure(4, weight=1)

    # ── Reports tab ───────────────────────────────────────────────────────────

    def _build_reports_tab(self):
        self.tab_reports = ctk.CTkFrame(self.content)
        top_bar = ctk.CTkFrame(self.tab_reports, fg_color="transparent")
        top_bar.pack(fill="x", padx=20, pady=15)

        ctk.CTkLabel(
            top_bar, text="Saved Patient Consultation Records",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(side="left")

        ctk.CTkButton(
            top_bar, text="📂 Open Records Folder",
            command=lambda: self._open_folder(config.OUTPUT_DIR),
            width=180, fg_color="#8A1538", hover_color="#6B102C"
        ).pack(side="right")

        self.reports_list = ctk.CTkTextbox(self.tab_reports, font=ctk.CTkFont(size=13))
        self.reports_list.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        self._refresh_reports_list()

    # ══════════════════════════════════════════════════════════════════════════
    # Navigation
    # ══════════════════════════════════════════════════════════════════════════

    def _show_tab(self, key: str):
        for w in [self.tab_process, self.tab_enroll, self.tab_reports]:
            w.pack_forget()
        if key == "process":
            self.tab_process.pack(fill="both", expand=True)
        elif key == "enroll":
            self.tab_enroll.pack(fill="both", expand=True)
        else:
            self.tab_reports.pack(fill="both", expand=True)
            self._refresh_reports_list()

    # ══════════════════════════════════════════════════════════════════════════
    # Backend Init
    # ══════════════════════════════════════════════════════════════════════════

    def _init_backend(self):
        try:
            self._log("Loading local AI models (Whisper STT, SpeechBrain, GGUF LLM)...")
            self.pipeline = AudioPipeline(progress_callback=self._log)
            self.report_gen = ReportGenerator()
            self._set_status(f"Ready | Offline Device: {config.DEVICE}")
            self._log("⚡ Local Multilingual AI Engine ready.")
        except Exception as e:
            self._set_status(f"Model Error: {e}")
            messagebox.showerror("Model Load Failed", str(e))

    # ══════════════════════════════════════════════════════════════════════════
    # Real-Time Live Recording & Transcription
    # ══════════════════════════════════════════════════════════════════════════

    def _toggle_recording(self):
        if not self.is_recording:
            self._start_recording()
        else:
            self._stop_recording()

    def _start_recording(self):
        if not self.pipeline:
            messagebox.showerror("Error", "AI models are still loading into memory. Please wait a moment.")
            return

        patient_name = self.entry_patient.get().strip()
        if not patient_name:
            messagebox.showerror("Required", "Please enter Patient Name before starting consultation.")
            return

        doc_name = self.entry_doctor.get().strip() or "Attending Physician"
        self.current_patient = patient_name
        self.current_attending = doc_name
        self.current_language = "ar" if "Arabic" in self.combo_language.get() else "en"

        # Reset state & UI
        self.current_transcript = []
        self.current_report = None
        self.report_box.delete("1.0", "end")
        self.report_box.insert("1.0", "🎙 Live Consultation in progress...\n\nDoctor and Patient are speaking. The AI is transcribing audio in real time in the background.\n\nClick 'Stop Consultation' when the visit concludes to generate the final physician-grade clinical report.")

        # Prepare RealtimeTranscriber reusing pipeline's models
        self.live_transcriber = RealtimeTranscriber(
            whisper_model=self.pipeline.whisper,
            embedder=self.pipeline.embedder,
            temp_dir=config.TEMP_DIR,
            doctor_name=doc_name
        )

        while not self._audio_queue.empty():
            try:
                self._audio_queue.get_nowait()
            except Exception:
                break

        self.is_recording = True
        self.btn_record.configure(text="⏹ Stop Consultation", fg_color="#e67e22", hover_color="#d35400")
        self.recording_status.configure(
            text="● Recording & Transcribing live in background... (Arabic, English, Hindi, Urdu, etc. supported)",
            text_color="#e74c3c"
        )
        self._log(f"🎙 Live consultation started for patient: {patient_name} (Doctor: {doc_name})")

        # Audio stream callback
        def _audio_callback(indata, frames, time_info, status):
            if status:
                pass
            if self.is_recording:
                self._audio_queue.put(indata.copy())

        self._recording_stream = sd.InputStream(
            samplerate=self.sample_rate,
            channels=1,
            dtype="float32",
            callback=_audio_callback,
        )
        self._recording_stream.start()

        # Start live transcription worker thread
        threading.Thread(target=self._live_transcription_worker, daemon=True).start()

    def _live_transcription_worker(self):
        """Worker thread: gathers audio in real-time slices."""
        chunk_target_samples = int(self.sample_rate * 6.0)
        accumulated_chunks = []
        current_samples = 0
        total_time_offset_s = 0.0
        profiles = self.pipeline.load_doctor_profiles()

        while self.is_recording or not self._audio_queue.empty():
            try:
                data = self._audio_queue.get(timeout=0.3)
                accumulated_chunks.append(data)
                current_samples += len(data)
            except queue.Empty:
                data = None

            if (current_samples >= chunk_target_samples) or (not self.is_recording and current_samples >= int(self.sample_rate * 0.8)):
                if accumulated_chunks:
                    audio_block = np.concatenate(accumulated_chunks, axis=0).squeeze()
                    accumulated_chunks = []
                    current_samples = 0
                    duration_s = len(audio_block) / self.sample_rate

                    try:
                        segs = self.live_transcriber.process_chunk(
                            audio_np=audio_block,
                            sample_rate=self.sample_rate,
                            profiles=profiles,
                            offset_s=total_time_offset_s,
                        )
                        total_time_offset_s += duration_s

                        if segs:
                            self.current_transcript.extend(segs)
                    except Exception as e:
                        pass

        # Live transcription is now complete!
        self.after(0, self._on_live_recording_finished)

    def _stop_recording(self):
        if not self.is_recording:
            return

        self.is_recording = False
        if self._recording_stream:
            self._recording_stream.stop()
            self._recording_stream.close()
            self._recording_stream = None

        self.btn_record.configure(text="⏳ Generating Report...", state="disabled", fg_color="#7f8c8d")
        self.recording_status.configure(text="● Synthesizing hospital-grade clinical consultation note with local LLM...", text_color="#2980b9")
        self._log("⏹ Consultation ended. Extracting clinical findings & physician directives...")

    def _on_live_recording_finished(self):
        threading.Thread(target=self._generate_report_thread, daemon=True).start()

    def _generate_report_thread(self):
        try:
            self.progress.start()
            if not self.current_transcript:
                self._log("Warning: No speech detected in recording.")
                messagebox.showwarning("No Speech", "No audible speech was detected during the recording.")
                self.btn_record.configure(text="🎙 Start Consultation", state="normal", fg_color="#8A1538")
                return

            self._log("⚡ Running semantic speaker verification & clinical note synthesis...")
            self.current_transcript = disambiguate_speakers(self.current_transcript, self.current_attending)

            report = self.report_gen.extract(
                self.current_transcript,
                doctor_name=self.current_attending,
                language=self.current_language
            )
            self.current_report = report

            text = self._render_text(report, self.current_patient, self.current_attending, self.current_language)
            
            def _update_ui():
                self.report_box.delete("1.0", "end")
                self.report_box.insert("1.0", text)
                self.recording_status.configure(text="✓ Clinical Consultation Note Generated!", text_color="#27ae60")
                self.btn_record.configure(text="🎙 Start Consultation", state="normal", fg_color="#8A1538")

            self.after(0, _update_ui)

            # Save JSON
            timestamp_str = time.strftime("%Y%m%d_%H%M%S")
            clean_name = self.current_patient.strip().replace(" ", "_")
            out_json = config.OUTPUT_DIR / f"Consultation_{clean_name}_{timestamp_str}_report.json"
            out_json.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
            self._log(f"✓ Saved clinical record JSON: {out_json.name}")
            self._refresh_reports_list()

        except Exception as e:
            self._log(f"ERROR generating report: {e}")
            messagebox.showerror("Report Error", str(e))
            self.recording_status.configure(text="Report generation failed", text_color="red")
            self.btn_record.configure(text="🎙 Start Consultation", state="normal", fg_color="#8A1538")
        finally:
            self.progress.stop()

    # ══════════════════════════════════════════════════════════════════════════
    # Audio File Pipeline
    # ══════════════════════════════════════════════════════════════════════════

    def _browse_audio(self):
        f = filedialog.askopenfilename(filetypes=[("Audio files", "*.wav *.mp3 *.m4a *.flac *.ogg")])
        if f:
            self.entry_audio.delete(0, "end")
            self.entry_audio.insert(0, f)

    def _run_file_pipeline(self):
        if not self.pipeline or not self.report_gen:
            messagebox.showinfo("Models Loading", "AI models are currently loading into memory. Please wait a moment until the status shows 'Ready'.")
            return

        path = self.entry_audio.get().strip()
        if not path or not Path(path).exists():
            messagebox.showerror("Error", "Please select a valid audio file.")
            return

        patient_name = self.entry_patient.get().strip() or "N/A"
        doctor_name = self.entry_doctor.get().strip() or "Attending Physician"
        fast_mode = bool(self.chk_fast_mode.get())
        self.current_language = "ar" if "Arabic" in self.combo_language.get() else "en"

        self.progress.start()
        threading.Thread(
            target=self._process_file_thread,
            args=(path, patient_name, doctor_name, fast_mode, self.current_language),
            daemon=True
        ).start()

    def _process_file_thread(self, audio_path: str, patient_name: str, doctor_name: str, fast_mode: bool, lang: str):
        try:
            if not self.pipeline or not self.report_gen:
                self._log("Waiting for AI models to finish loading...")
                while not self.pipeline or not self.report_gen:
                    time.sleep(0.5)

            self.report_box.delete("1.0", "end")
            self.report_box.insert("1.0", f"Processing consultation audio file: {Path(audio_path).name}...\n\nExtracting clinical findings and transcribing with Whisper...")

            self.current_patient = patient_name
            self.current_attending = doctor_name
            
            # Safely load doctor profiles
            profiles = self.pipeline.load_doctor_profiles() if self.pipeline else {}

            if fast_mode:
                self._log(f"⚡ Fast Processing: {Path(audio_path).name}...")
                waveform, sr = load_audio(audio_path)
                transcriber = RealtimeTranscriber(
                    whisper_model=self.pipeline.whisper,
                    embedder=self.pipeline.embedder,
                    temp_dir=config.TEMP_DIR,
                    doctor_name=doctor_name
                )
                labeled = transcriber.process_file_fast(
                    audio_path=audio_path,
                    waveform=waveform,
                    sample_rate=sr,
                    profiles=profiles,
                    progress_cb=self._log
                )
            else:
                self._log(f"Running Full Diarization on: {Path(audio_path).name}...")
                result = self.pipeline.process(audio_path, patient_name)
                labeled = result["labeled_transcript"]

            # Semantic Role Verification
            labeled = disambiguate_speakers(labeled, doctor_name)
            self.current_transcript = labeled

            wc = sum(len(s["text"].split()) for s in labeled)
            self._log(f"✓ Speech transcribed: {len(labeled)} segments, ~{wc} words. Synthesizing clinical note...")

            # Generate Report
            report = self.report_gen.extract(labeled, doctor_name=doctor_name, language=lang)
            self.current_report = report

            text = self._render_text(report, patient_name, doctor_name, lang)

            def _show_rp():
                self.report_box.delete("1.0", "end")
                self.report_box.insert("1.0", text)
            self.after(0, _show_rp)

            # Save JSON
            stem = Path(audio_path).stem
            out_json = config.OUTPUT_DIR / f"{stem}_report.json"
            out_json.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
            self._log(f"✓ Saved clinical record: {out_json.name}")
            self._refresh_reports_list()

        except Exception as e:
            self._log(f"ERROR in file pipeline: {e}")
            messagebox.showerror("Processing Error", str(e))
        finally:
            self.progress.stop()

    # ── Dialogue Transcript Modal / Popup ─────────────────────────────────────

    def _toggle_transcript_dialog(self):
        """Shows the raw transcribed dialogue in a clean secondary window for review."""
        if not self.current_transcript:
            messagebox.showinfo("Transcript", "No consultation dialogue recorded yet.")
            return

        top = ctk.CTkToplevel(self)
        top.title("Consultation Dialogue Transcript (Auditing View)")
        top.geometry("850x650")
        top.attributes('-topmost', True)

        ctk.CTkLabel(
            top, text="Raw Consultation Dialogue Transcript",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(pady=10)

        tx_box = ctk.CTkTextbox(top, font=ctk.CTkFont(size=13))
        tx_box.pack(fill="both", expand=True, padx=15, pady=10)

        for seg in self.current_transcript:
            ts = f"[{seg['start']:.1f}s]"
            role = seg["role"]
            name = seg.get("doctor_name") or ""
            speaker = role + (f" ({name})" if name and role == "Doctor" else "")
            tx_box.insert("end", f"{ts} {speaker}:\n  {seg['text']}\n\n")

        btn_frame = ctk.CTkFrame(top, fg_color="transparent")
        btn_frame.pack(fill="x", padx=15, pady=10)

        def _copy_tx():
            lines = [f"[{s['start']:.1f}s] {s['role']}: {s['text']}" for s in self.current_transcript]
            self.clipboard_clear()
            self.clipboard_append("\n".join(lines))
            self._log("✓ Transcript copied to clipboard.")
            top.destroy()

        ctk.CTkButton(btn_frame, text="📋 Copy Full Transcript", command=_copy_tx, width=180).pack(side="left")
        ctk.CTkButton(btn_frame, text="Close", command=top.destroy, width=100).pack(side="right")

    def _copy_report(self):
        text = self.report_box.get("1.0", "end").strip()
        if not text:
            return
        self.clipboard_clear()
        self.clipboard_append(text)
        self._log("✓ Clinical consultation note copied to clipboard.")

    # ══════════════════════════════════════════════════════════════════════════
    # Doctor Enrollment
    # ══════════════════════════════════════════════════════════════════════════

    def _browse_enroll(self):
        f = filedialog.askopenfilename(filetypes=[("Audio", "*.wav *.mp3 *.m4a *.flac")])
        if f:
            self.entry_enroll_path = f
            self.enroll_log.insert("end", f"Selected audio file: {f}\n")

    def _run_enroll(self):
        name = self.entry_doc_name.get().strip()
        if not name or not hasattr(self, "entry_enroll_path"):
            messagebox.showerror("Required", "Please enter Doctor Name and select an audio sample.")
            return
        threading.Thread(
            target=self._enroll_thread, args=(name, self.entry_enroll_path), daemon=True
        ).start()

    def _enroll_thread(self, name: str, path: str):
        try:
            safe = self.pipeline.enroll_doctor(name, path)
            self.enroll_log.insert("end", f"✓ Successfully Enrolled: {name} (ID: {safe})\n")
            self._refresh_reports_list()
        except Exception as e:
            self.enroll_log.insert("end", f"ERROR during enrollment: {e}\n")

    # ══════════════════════════════════════════════════════════════════════════
    # Hospital-Grade Word DOCX Export
    # ══════════════════════════════════════════════════════════════════════════

    def _export_docx(self):
        if not self.current_report:
            messagebox.showwarning("Warning", "No report available to export.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not path:
            return

        from datetime import datetime
        doc = Document()

        # Page Setup
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Header Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("CLINICAL CONSULTATION REPORT")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(138, 21, 56)  # Maroon

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run("Department of Clinical Medicine • State of Qatar")
        sub_run.font.name = "Arial"
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(100, 100, 100)

        # Demographics Table
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.style = "Table Grid"

        hdr_cells = meta_table.rows[0].cells
        hdr_cells[0].text = f"Patient Name: {self.current_patient}"
        hdr_cells[1].text = f"Date / Time: {datetime.now().strftime('%Y-%m-%d %H:%M')}"

        row1_cells = meta_table.rows[1].cells
        row1_cells[0].text = f"Attending Physician: {self.current_attending}"
        row1_cells[1].text = "Clinical Facility: Healthcare Center, Doha"

        row2_cells = meta_table.rows[2].cells
        row2_cells[0].text = "Encounter Type: In-Person Clinical Consultation"
        row2_cells[1].text = "Documentation: Board-Certified Scribe AI"

        doc.add_paragraph().paragraph_format.space_after = 10

        r = self.current_report
        is_ar = (self.current_language == "ar")

        # 1. Chief Complaint
        h1 = doc.add_heading("1. Chief Complaint" if not is_ar else "١. الشكوى الرئيسية", level=2)
        doc.add_paragraph(r.get("chief_complaint") or ("Not specified." if not is_ar else "غير محدد."))

        # 2. HPI / Subjective Notes
        h2 = doc.add_heading("2. History of Present Illness (Subjective)" if not is_ar else "٢. تاريخ المرض الحالي", level=2)
        doc.add_paragraph(r.get("history_notes") or ("No history details recorded." if not is_ar else "لا توجد تفاصيل مسجلة."))

        # 3. Physical Examination
        h3 = doc.add_heading("3. Physical Examination & Clinical Findings (Objective)" if not is_ar else "٣. الفحص السريري والملاحظات الموضوعية", level=2)
        doc.add_paragraph(r.get("examination_findings") or ("Deferred / Not conducted during this consultation." if not is_ar else "تم تأجيله / لم يُجرَ خلال هذه الجلسة."))

        # 4. Assessment & Diagnosis
        h4 = doc.add_heading("4. Clinical Assessment & Working Diagnosis" if not is_ar else "٤. التقييم والتشخيص السريري", level=2)
        doc.add_paragraph(r.get("diagnosis") or ("Pending clinical review." if not is_ar else "قيد المراجعة السريرية."))

        # 5. Prescriptions
        h5 = doc.add_heading("5. Pharmacotherapy & Prescriptions" if not is_ar else "٥. الخطة العلاجية والأدوية الموصوفة", level=2)
        prescriptions = [p for p in (r.get("prescriptions") or []) if p.get("medication")]
        if prescriptions:
            rx_table = doc.add_table(rows=1, cols=3)
            rx_table.style = "Table Grid"
            hdr = rx_table.rows[0].cells
            hdr[0].text = "Medication" if not is_ar else "الدواء"
            hdr[1].text = "Dosage / Strength" if not is_ar else "الجرعة"
            hdr[2].text = "Instructions & Frequency" if not is_ar else "طريقة الاستخدام والتكرار"
            for item in prescriptions:
                row = rx_table.add_row().cells
                row[0].text = str(item.get("medication") or "—")
                row[1].text = str(item.get("dosage") or "—")
                row[2].text = str(item.get("instructions") or "—")
        else:
            doc.add_paragraph("No new pharmacological prescriptions ordered." if not is_ar else "لا توجد أدوية جديدة موصوفة في هذه الزيارة.")

        # 6. Diagnostic Orders
        h6 = doc.add_heading("6. Diagnostic Investigations & Orders" if not is_ar else "٦. الفحوصات التشخيصية المطلوبة", level=2)
        tests = [t for t in (r.get("tests_ordered") or []) if t and str(t).strip()]
        if tests:
            for test in tests:
                doc.add_paragraph(str(test), style="List Bullet")
        else:
            doc.add_paragraph("No additional diagnostic laboratory or radiological investigations ordered." if not is_ar else "لا توجد فحوصات مخبرية أو إشعاعية إضافية مطلوبة.")

        # 7. Follow-up
        h7 = doc.add_heading("7. Follow-up & Red-Flag Precautions" if not is_ar else "٧. خطة المتابعة وتحذيرات الطوارئ", level=2)
        doc.add_paragraph(r.get("follow_up") or ("Routine clinical follow-up as needed." if not is_ar else "متابعة سريرية روتينية حسب الحاجة."))

        # 8. Patient Instructions
        if r.get("other_instructions"):
            h8 = doc.add_heading("8. Patient Education & Care Instructions" if not is_ar else "٨. إرشادات وتثقيف المريض", level=2)
            doc.add_paragraph(str(r.get("other_instructions")))

        # Signature
        doc.add_paragraph().paragraph_format.space_after = 20
        sig_p = doc.add_paragraph()
        sig_p.add_run(f"\nAttending Physician: {self.current_attending}\nElectronic Verification Record • Confidential Medical Record").font.size = Pt(9)

        doc.save(path)
        self._log(f"✓ Exported Word Document: {path}")
        messagebox.showinfo("Export Successful", f"Clinical consultation note saved to:\n{path}")

    # ══════════════════════════════════════════════════════════════════════════
    # Formatting Helpers
    # ══════════════════════════════════════════════════════════════════════════

    def _render_text(self, r, patient, doctor, lang="en") -> str:
        is_ar = (lang == "ar")
        if is_ar:
            lines = [
                "=" * 70,
                "                   تقرير الاستشارة الطبية السريرية",
                "=" * 70,
                f"اسم المريض:       {patient}",
                f"الطبيب المعالج:   {doctor}",
                f"المنشأة الصحية:   مركز الرعاية الصحية • دولة قطر",
                "-" * 70,
                "١. الشكوى الرئيسية (Chief Complaint):",
                f"   {r.get('chief_complaint') or 'غير محدد'}",
                "",
                "٢. تاريخ المرض الحالي (History of Present Illness - HPI):",
                f"   {r.get('history_notes') or 'لا توجد تفاصيل مسجلة'}",
                "",
                "٣. الفحص السريري والملاحظات (Physical Examination & Objective Findings):",
                f"   {r.get('examination_findings') or 'تم تأجيله / لم يُجرَ خلال هذه الجلسة'}",
                "",
                "٤. التقييم والتشخيص السريري (Assessment & Working Diagnosis):",
                f"   {r.get('diagnosis') or 'قيد المراجعة السريرية'}",
                "",
                "٥. الفحوصات التشخيصية المطلوبة (Diagnostic Orders):",
            ]
            tests = [t for t in (r.get("tests_ordered") or []) if t and str(t).strip()]
            if tests:
                lines += [f"   • {t}" for t in tests]
            else:
                lines.append("   • لا توجد فحوصات مخبرية أو إشعاعية مطلوبة")

            lines += ["", "٦. الخطة العلاجية والأدوية الموصوفة (Prescriptions):"]
            prescriptions = [p for p in (r.get("prescriptions") or []) if p.get("medication")]
            if prescriptions:
                for p in prescriptions:
                    med_line = f"   • {p.get('medication')}"
                    if p.get("dosage"):
                        med_line += f" | الجرعة: {p.get('dosage')}"
                    if p.get("instructions"):
                        med_line += f" | طريقة الاستخدام: {p.get('instructions')}"
                    lines.append(med_line)
            else:
                lines.append("   • لا توجد أدوية جديدة موصوفة في هذه الزيارة")

            lines += [
                "",
                "٧. خطة المتابعة وتحذيرات الطوارئ (Follow-up & Care Plan):",
                f"   {r.get('follow_up') or 'متابعة روتينية حسب الحاجة'}",
            ]
            if r.get("other_instructions"):
                lines += ["", "٨. إرشادات وتثقيف المريض (Patient Care Instructions):", f"   {r.get('other_instructions')}"]

            lines.append("=" * 70)
            return "\n".join(lines)

        # Standard Medical English Note (SOAP format)
        lines = [
            "=" * 70,
            "                   CLINICAL CONSULTATION REPORT",
            "         Department of Clinical Medicine • State of Qatar",
            "=" * 70,
            f"Patient Name:        {patient}",
            f"Attending Physician: {doctor}",
            f"Encounter Location:  Clinical Center, Doha • State of Qatar",
            "-" * 70,
            "I. SUBJECTIVE",
            "   CHIEF COMPLAINT:",
            f"     {r.get('chief_complaint') or '—'}",
            "",
            "   HISTORY OF PRESENT ILLNESS (HPI):",
            f"     {r.get('history_notes') or '—'}",
            "",
            "II. OBJECTIVE",
            "   PHYSICAL EXAMINATION & VITAL SIGNS:",
            f"     {r.get('examination_findings') or 'Deferred / Not conducted during this consultation'}",
            "",
            "III. ASSESSMENT",
            "   CLINICAL IMPRESSION & WORKING DIAGNOSIS:",
            f"     {r.get('diagnosis') or 'Clinical assessment pending diagnostic review'}",
            "",
            "IV. PLAN & MANAGEMENT",
            "   DIAGNOSTIC ORDERS & INVESTIGATIONS:",
        ]
        tests = [t for t in (r.get("tests_ordered") or []) if t and str(t).strip()]
        if tests:
            lines += [f"     • {t}" for t in tests]
        else:
            lines.append("     • None ordered")

        lines += ["", "   PHARMACOTHERAPY & PRESCRIPTIONS:"]
        prescriptions = [p for p in (r.get("prescriptions") or []) if p.get("medication")]
        if prescriptions:
            for p in prescriptions:
                med_line = f"     • {p.get('medication')}"
                if p.get("dosage"):
                    med_line += f" | Dose: {p.get('dosage')}"
                if p.get("instructions"):
                    med_line += f" | Instructions: {p.get('instructions')}"
                lines.append(med_line)
        else:
            lines.append("     • None prescribed")

        lines += [
            "",
            "   FOLLOW-UP & RED-FLAG PRECAUTIONS:",
            f"     {r.get('follow_up') or 'Routine follow-up as needed'}",
        ]
        if r.get("other_instructions"):
            lines += ["", "   PATIENT EDUCATION & CARE INSTRUCTIONS:", f"     {r.get('other_instructions')}"]

        lines.append("=" * 70)
        return "\n".join(lines)

    def _set_placeholder_report(self):
        self.report_box.delete("1.0", "end")
        placeholder = """======================================================================
                   CLINICAL CONSULTATION REPORT
         Department of Clinical Medicine • State of Qatar
======================================================================
Patient Name:        [Enter patient name above]
Attending Physician: [Enter doctor name above]
Clinical Facility:   Healthcare Center, Doha • State of Qatar
----------------------------------------------------------------------
I. SUBJECTIVE
   CHIEF COMPLAINT:
     [Consultation note will be generated automatically when visit ends]

   HISTORY OF PRESENT ILLNESS (HPI):
     [Detailed narrative of onset, duration, symptoms, and medical history]

II. OBJECTIVE
   PHYSICAL EXAMINATION & VITAL SIGNS:
     [Observations, clinical signs, and physical examination findings]

III. ASSESSMENT
   CLINICAL IMPRESSION & WORKING DIAGNOSIS:
     [Working clinical diagnosis and differential diagnoses]

IV. PLAN & MANAGEMENT
   DIAGNOSTIC ORDERS & INVESTIGATIONS:
     • [Laboratory and radiological imaging orders]

   PHARMACOTHERAPY & PRESCRIPTIONS:
     • [Prescribed medications, dosage, frequency, and instructions]

   FOLLOW-UP & RED-FLAG PRECAUTIONS:
     [Follow-up schedule and emergency return precautions]
======================================================================"""
        self.report_box.insert("1.0", placeholder)

    def _refresh_reports_list(self):
        self.reports_list.delete("1.0", "end")
        files = sorted(config.OUTPUT_DIR.glob("*_report.json"), reverse=True)
        self.reports_list.insert("1.0", f"Saved Consultation Records ({len(files)} total):\n\n")
        for f in files:
            self.reports_list.insert("end", f"📄 {f.name}\n")

    def _open_folder(self, folder: Path):
        import platform
        import subprocess
        p = platform.system()
        if p == "Windows":
            subprocess.run(["explorer", str(folder)])
        elif p == "Darwin":
            subprocess.run(["open", str(folder)])
        else:
            subprocess.run(["xdg-open", str(folder)])

    def _log(self, msg: str):
        self.log_label.configure(text=msg)

    def _set_status(self, text: str):
        self.status.configure(text=text)


def main():
    app = MedicalScribeApp()
    app.mainloop()
